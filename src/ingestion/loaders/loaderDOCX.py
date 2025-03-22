from src.ingestion.loaders.loaderBase import LoaderBase
from docx import Document

class LoaderDOCX(LoaderBase):

    def __init__(self, filepath:str):
        self.filepath=filepath

    def extract_metadata(self):
        docx_file_reader = Document(self.filepath)
        doc_info = docx_file_reader.core_properties
        metadata = {  
            'author': doc_info.author,  
            # 'creator': doc_info.creator,  
            # 'producer': doc_info.producer,  
            'subject': doc_info.subject,  
            'title': doc_info.title,  
            #'number_of_pages': len(doc_info.pages)  
        }

        self.metadata=metadata
        
        return self.metadata if self.all_keys_have_values(metadata=self.metadata) else False
    
    def extract_text(self):
        document = Document(self.filepath)
        text = "\n".join([p.text for p in document.paragraphs])
        return text